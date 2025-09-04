"""Resume parser for PDF and DOCX files."""

import re
from pathlib import Path

import pdfplumber
from docx import Document as DocxDocument

from ..database.models import ResumeData


class ResumeParser:
    """Parse resumes from PDF and DOCX files."""

    def __init__(self):
        """Initialize parser with common patterns."""
        self.email_pattern = re.compile(
            r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b"
        )
        self.phone_pattern = re.compile(
            r"(?:\+?1[-.\s]?)?\(?(\d{3})\)?[-.\s]?(\d{3})[-.\s]?(\d{4})"
        )
        self.skill_keywords = {
            "programming": [
                "Python", "Java", "JavaScript", "TypeScript", "C++", "C#",
                "Ruby", "Go", "Rust", "Swift", "Kotlin", "PHP", "Scala",
            ],
            "web": [
                "React", "Angular", "Vue", "Node.js", "Django", "Flask",
                "FastAPI", "Express", "Spring", "ASP.NET", "Rails",
            ],
            "database": [
                "SQL", "PostgreSQL", "MySQL", "MongoDB", "Redis", "Cassandra",
                "DynamoDB", "Elasticsearch", "SQLite", "Oracle",
            ],
            "cloud": [
                "AWS", "Azure", "GCP", "Docker", "Kubernetes", "Terraform",
                "Jenkins", "CI/CD", "DevOps", "CloudFormation",
            ],
            "data": [
                "Pandas", "NumPy", "Scikit-learn", "TensorFlow", "PyTorch",
                "Spark", "Hadoop", "Tableau", "Power BI", "Machine Learning",
            ],
        }

    def parse(self, file_path: str) -> ResumeData:
        """Parse resume from file path."""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Resume file not found: {file_path}")

        if path.suffix.lower() == ".pdf":
            text = self._extract_pdf(file_path)
        elif path.suffix.lower() in [".docx", ".doc"]:
            text = self._extract_docx(file_path)
        elif path.suffix.lower() == ".txt":
            text = self._extract_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {path.suffix}")

        # Extract name before cleaning (need newlines for name extraction)
        name = self._extract_name(text)
        
        # Clean text for other extractions
        text = self._clean_text(text)

        # Extract components
        return ResumeData(
            name=name,
            email=self._extract_email(text) or "unknown@example.com",
            phone=self._extract_phone(text),
            location=self._extract_location(text),
            skills=self._extract_skills(text),
            experience=self._extract_experience(text),
            education=self._extract_education(text),
            preferred_roles=self._extract_preferred_roles(text),
        )

    def _extract_pdf(self, path: str) -> str:
        """Extract text from PDF file using pdfplumber."""
        text = ""
        try:
            with pdfplumber.open(path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"

                    # Extract tables
                    tables = page.extract_tables()
                    for table in tables:
                        if table:
                            text += self._table_to_text(table) + "\n"
        except Exception as e:
            raise ValueError(f"Failed to parse PDF: {e}") from e

        return text

    def _extract_docx(self, path: str) -> str:
        """Extract text from DOCX file."""
        text = ""
        try:
            doc = DocxDocument(path)

            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text += row_text + "\n"
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX: {e}") from e

        return text

    def _extract_txt(self, path: str) -> str:
        """Extract text from TXT file."""
        try:
            with open(path, 'r', encoding='utf-8') as f:
                text = f.read()
        except Exception as e:
            raise ValueError(f"Failed to parse TXT: {e}") from e
        return text

    def _table_to_text(self, table: list[list[str]]) -> str:
        """Convert table to text format."""
        if not table:
            return ""

        text_rows = []
        for row in table:
            if row and any(cell for cell in row if cell):
                text_rows.append(" | ".join(str(cell or "") for cell in row))

        return "\n".join(text_rows)

    def _clean_text(self, text: str) -> str:
        """Clean and normalize text."""
        # Remove excessive whitespace
        text = re.sub(r"\s+", " ", text)
        # Remove special characters but keep basic punctuation
        text = re.sub(r"[^\w\s@.-]", " ", text)
        return text.strip()

    def _extract_name(self, text: str) -> str:
        """Extract name from resume text."""
        # Usually the name is at the beginning
        lines = text.split("\n")
        for line in lines[:5]:  # Check first 5 lines
            # Skip lines with email or phone
            if "@" in line or any(char.isdigit() for char in line):
                continue
            # Look for capitalized words
            words = line.strip().split()
            if 2 <= len(words) <= 4:
                if all(word[0].isupper() for word in words if word):
                    return " ".join(words)

        return "Unknown"

    def _extract_email(self, text: str) -> str | None:
        """Extract email from resume text."""
        match = self.email_pattern.search(text)
        return match.group(0) if match else None

    def _extract_phone(self, text: str) -> str | None:
        """Extract phone number from resume text."""
        match = self.phone_pattern.search(text)
        if match:
            groups = match.groups()
            return f"({groups[0]}) {groups[1]}-{groups[2]}"
        return None

    def _extract_location(self, text: str) -> str:
        """Extract location from resume text."""
        # Common location patterns
        location_patterns = [
            r"(?:Location|Address|City)[:]\s*([^,\n]+(?:,\s*[A-Z]{2})?)",
            r"([A-Za-z\s]+,\s*[A-Z]{2}\s*\d{5})",  # City, STATE ZIP
            r"([A-Za-z\s]+,\s*[A-Z]{2})",  # City, STATE
        ]

        for pattern in location_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(1).strip()

        return "Not specified"

    def _extract_skills(self, text: str) -> list[str]:
        """Extract skills from resume text."""
        skills = set()
        text.lower()

        # Look for skills section
        skills_section = self._extract_section(text, ["skills", "technical skills", "core competencies"])
        if skills_section:
            text_to_search = skills_section
        else:
            text_to_search = text

        # Search for known skills
        for _category, skill_list in self.skill_keywords.items():
            for skill in skill_list:
                if skill.lower() in text_to_search.lower():
                    skills.add(skill)

        # Also look for skills in comma-separated lists
        skill_pattern = r"(?:Skills|Technologies|Tools)[:]\s*([^.\n]+)"
        match = re.search(skill_pattern, text, re.IGNORECASE)
        if match:
            skill_text = match.group(1)
            for skill in re.split(r"[,;|]", skill_text):
                skill = skill.strip()
                if skill and len(skill) < 30:  # Reasonable skill name length
                    skills.add(skill)

        return list(skills)

    def _extract_experience(self, text: str) -> list[dict[str, str]]:
        """Extract work experience from resume text."""
        experience = []

        # Look for experience section
        exp_section = self._extract_section(
            text, ["experience", "work experience", "employment", "professional experience"]
        )

        if not exp_section:
            return experience

        # Pattern for job entries (Title at Company, Date)
        job_pattern = r"([A-Za-z\s]+)\s+(?:at|@|-)\s+([A-Za-z\s&.]+)[,\s]+(\d{4}(?:\s*-\s*\d{4}|\s*-\s*Present)?)"
        matches = re.finditer(job_pattern, exp_section, re.IGNORECASE)

        for match in matches:
            experience.append({
                "title": match.group(1).strip(),
                "company": match.group(2).strip(),
                "duration": match.group(3).strip(),
                "description": "",  # Would need more complex parsing for descriptions
            })

        return experience[:5]  # Return top 5 experiences

    def _extract_education(self, text: str) -> list[dict[str, str]]:
        """Extract education from resume text."""
        education = []

        # Look for education section
        edu_section = self._extract_section(text, ["education", "academic", "qualifications"])

        if not edu_section:
            return education

        # Pattern for degree entries
        degree_patterns = [
            r"(Bachelor|Master|PhD|Ph\.D\.|B\.S\.|M\.S\.|MBA|B\.A\.|M\.A\.)[^,\n]*(?:in\s+)?([^,\n]+)",
            r"([A-Za-z\s]+University|College|Institute)[^,\n]*",
        ]

        for pattern in degree_patterns:
            matches = re.finditer(pattern, edu_section, re.IGNORECASE)
            for match in matches:
                if len(education) < 3:  # Limit to 3 education entries
                    education.append({
                        "degree": match.group(0).strip(),
                        "school": "",
                        "year": self._extract_year(match.group(0)),
                    })

        return education

    def _extract_preferred_roles(self, text: str) -> list[str]:
        """Extract preferred job roles from resume text."""
        roles = []

        # Common role titles
        role_keywords = [
            "Software Engineer", "Full Stack Developer", "Backend Engineer",
            "Frontend Developer", "Data Scientist", "DevOps Engineer",
            "Product Manager", "Technical Lead", "Solutions Architect",
            "Machine Learning Engineer", "Cloud Engineer", "Site Reliability Engineer",
        ]

        # Check objective/summary section
        objective = self._extract_section(text, ["objective", "summary", "profile"])

        text_to_search = objective if objective else text

        for role in role_keywords:
            if role.lower() in text_to_search.lower():
                roles.append(role)

        return roles[:3]  # Return top 3 preferred roles

    def _extract_section(self, text: str, keywords: list[str]) -> str | None:
        """Extract a section from resume based on keywords."""
        lines = text.split("\n")

        for i, line in enumerate(lines):
            line_lower = line.lower().strip()
            if any(keyword in line_lower for keyword in keywords):
                # Found section header, extract until next section or end
                section_lines = []
                for j in range(i + 1, min(i + 50, len(lines))):  # Limit section size
                    next_line = lines[j].strip()
                    # Check if this is a new section
                    if next_line and next_line[0].isupper() and ":" in next_line:
                        break
                    section_lines.append(next_line)

                return "\n".join(section_lines)

        return None

    def _extract_year(self, text: str) -> str:
        """Extract year from text."""
        year_match = re.search(r"\b(19|20)\d{2}\b", text)
        return year_match.group(0) if year_match else ""
