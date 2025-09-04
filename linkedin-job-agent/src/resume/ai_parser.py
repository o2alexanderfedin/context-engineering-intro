"""AI-powered resume parser using Claude Code SDK."""

import os
from pathlib import Path
from typing import Optional

from pdfplumber import PDF
from docx import Document as DocxDocument

from ..utils.json_utils import extract_json_from_text
from ..utils.claude_utils import ask_claude

class AIResumeParser:
    """Parse resume using AI for all extraction."""
    
    def __init__(self):
        """Initialize AI parser and load OAuth token."""
        # Load OAuth token from environment or config
        from ..config import get_settings
        settings = get_settings()
        
        # Set OAuth token if available
        if settings.claude_code_oauth_token:
            os.environ['CLAUDE_CODE_OAUTH_TOKEN'] = settings.claude_code_oauth_token
            del os.environ['ANTHROPIC_API_KEY']
        
        # Add ~/bin to PATH for claude CLI
        os.environ['PATH'] = os.path.expanduser('~/bin') + ':' + os.environ.get('PATH', '')

    async def parse(self, file_path: str) -> dict:
        """Parse resume using AI."""
        raw_text = self._extract_raw_text(file_path)

        # Now use AI to extract ALL information
        prompt = f"""
RESUME TEXT:
```markdown
{raw_text}
```
        
Analyze this resume and extract ALL information in JSON format.

Extract:
- name: Full name of the person
- email: Email address
- phone: Phone number
- location: Location/city
- summary: Professional summary
- skills: List of ALL technical and professional skills mentioned
- experience: List of work experiences with company, role, duration, and responsibilities
- education: Educational background
- preferred_roles: Based on the resume, what job roles would this person be suited for
- keywords: List of important keywords for job searching
- strengths: Key strengths and accomplishments
- raw_text: Full resume markdown text

Return ONLY valid JSON."""

        system_prompt = "You are an expert resume parser and job matching assistant. Extract structured information accurately and return valid JSON."
        result = await ask_claude(prompt, system_prompt=system_prompt, debug=True)
        
        try:
            # Use the global function to extract and parse JSON
            parsed = extract_json_from_text(result)
            
            # Ensure we have at least a name
            if not parsed.get('name'):
                raise ValueError(f"Failed to parse AI response: name not parsed")
            if not parsed.get('email'):
                raise ValueError(f"Failed to parse AI response: email not parsed")
            if not parsed.get('raw_text'):
                parsed['raw_text'] = raw_text
            return parsed
        except Exception as e:
            # If JSON parsing fails, raise an error with details
            raise ValueError(f"Failed to parse AI response: {str(e)}\nResponse was: {result[:500] if result else 'empty'}")
    
    async def get_keywords(self, resume_text: str) -> list[str]:
        """Get job search keywords from resume using AI."""
        prompt = f"""Extract job search keywords from this resume.

Focus on:
- Technical skills
- Programming languages
- Frameworks and tools
- Job titles
- Industry terms

Return as JSON array of keywords.

RESUME:
{resume_text}

Return ONLY JSON array like: ["keyword1", "keyword2", ...]"""

        result = await ask_claude(prompt, debug=False)
        
        try:
            # Use the global function to extract and parse JSON
            parsed = extract_json_from_text(result)
            # Ensure it's a list
            if isinstance(parsed, list):
                return parsed
            return []
        except:
            return []

    async def match_job(self, resume_text: str, job_description: str) -> dict:
        """Match resume to job using AI."""
        prompt = f"""Analyze how well this resume matches this job description.

Provide:
- match_score: 0-100 score
- matching_skills: Skills that match
- missing_skills: Required skills not in resume
- strengths: Why this is a good match
- weaknesses: Potential concerns
- recommendation: Should apply? (yes/no/maybe)

RESUME:
{resume_text}

JOB DESCRIPTION:
{job_description}

Return ONLY valid JSON."""

        result = await ask_claude(prompt, debug=False)
        
        if not result:
            return {"error": "No AI response", "match_score": 50, "recommendation": "maybe"}
        
        try:
            # Use the global function to extract and parse JSON
            parsed = extract_json_from_text(result)
            # Ensure required fields
            if 'match_score' not in parsed:
                parsed['match_score'] = 50
            if 'recommendation' not in parsed:
                parsed['recommendation'] = 'maybe'
            return parsed
        except:
            return {"error": "Failed to parse AI response", "match_score": 50, "recommendation": "maybe", "raw": result}

    def _extract_raw_text(self, file_path: str) -> str:
        """Extract raw text from file (PDF, DOCX, or TXT)."""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Resume file not found: {file_path}")

        if path.suffix.lower() == ".pdf":
            return self._extract_pdf(file_path)
        elif path.suffix.lower() in [".docx", ".doc"]:
            return self._extract_docx(file_path)
        elif path.suffix.lower() in [".txt", ".md"]:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        else:
            raise ValueError(f"Unsupported file format: {path.suffix}")

    def _extract_pdf(self, path: str) -> str:
        """Extract text from PDF file."""
        text = ""
        try:
            with PDF.open(path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            raise ValueError(f"Failed to parse PDF: {e}") from e
        return text

    def _extract_docx(self, path: str) -> str:
        """Extract text from DOCX file."""
        text = ""
        try:
            doc = DocxDocument(path)
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text += row_text + "\n"
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX: {e}") from e
        return text

    
